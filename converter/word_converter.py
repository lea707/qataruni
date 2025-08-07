import os
from docx import Document
from .base import DocumentConverter

class WordConverter(DocumentConverter):
    def load_document(self):
        if hasattr(self.input_path, 'read'):  # file-like object
            self.doc = Document(self.input_path)
        elif isinstance(self.input_path, str):
            if not os.path.exists(self.input_path):
                raise FileNotFoundError(f"❌ File not found: {self.input_path}")
            if not self.input_path.lower().endswith('.docx'):
                raise ValueError("❌ Only .docx files are supported.")
            self.doc = Document(self.input_path)
        else:
            raise TypeError("❌ input_path must be a file path or file-like object.")

    def convert_to_text(self):
        if not hasattr(self, 'doc'):
            raise RuntimeError("⚠️ Document not loaded. Call load_document() first.")
        
        self.text = "\n".join([
            para.text for para in self.doc.paragraphs if para.text.strip()
        ])